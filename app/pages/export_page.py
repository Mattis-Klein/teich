from __future__ import annotations

from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Optional, List, Dict, Any
import json
import time

from PySide6 import QtWidgets, QtCore

from ..widgets import Card, TopBar
from ..store import DataStore, Project


@dataclass
class ExportTemplate:
    id: str
    name: str
    rtl: bool = True
    include_header: bool = True
    font_size: int = 12


class ExportPage(QtWidgets.QWidget):
    back_to_editor = QtCore.Signal()

    def __init__(self, ds: DataStore, parent=None):
        super().__init__(parent)
        self.ds = ds
        self._project: Optional[Project] = None
        self._templates_path = self.ds.root / "templates.json"
        self._templates: List[ExportTemplate] = []

        outer = QtWidgets.QVBoxLayout(self)
        outer.setContentsMargins(18, 18, 18, 18)
        outer.setSpacing(12)

        self.top = TopBar("Export", show_back=True, show_home=False)
        self.top.back_clicked.connect(self.back_to_editor.emit)
        outer.addWidget(self.top)

        card = Card()
        lay = QtWidgets.QVBoxLayout(card)
        lay.setContentsMargins(16, 16, 16, 16)
        lay.setSpacing(12)

        self.lbl_title = QtWidgets.QLabel("—")
        self.lbl_title.setObjectName("h2")
        lay.addWidget(self.lbl_title)

        row = QtWidgets.QHBoxLayout()
        row.setSpacing(10)
        self.lst = QtWidgets.QListWidget()
        self.lst.setMinimumWidth(280)
        self.lst.setFixedHeight(240)

        right = QtWidgets.QVBoxLayout()
        self.preview = QtWidgets.QLabel("Pick a template to preview")
        self.preview.setWordWrap(True)
        self.preview.setMinimumHeight(120)
        self.preview.setObjectName("muted")

        self.btn_add = QtWidgets.QPushButton("Add template")
        self.btn_add.setFixedHeight(36)
        self.btn_add.setCursor(QtCore.Qt.PointingHandCursor)

        self.btn_export = QtWidgets.QPushButton("Export to .docx")
        self.btn_export.setFixedHeight(40)
        self.btn_export.setCursor(QtCore.Qt.PointingHandCursor)

        right.addWidget(self.preview)
        right.addStretch(1)
        right.addWidget(self.btn_add)
        right.addWidget(self.btn_export)

        row.addWidget(self.lst)
        row.addLayout(right, 1)
        lay.addLayout(row)

        outer.addWidget(card)
        outer.addStretch(1)

        self.lst.currentRowChanged.connect(self._render_preview)
        self.btn_add.clicked.connect(self._add_template)
        self.btn_export.clicked.connect(self._do_export)

        self._load_templates()

    def set_project(self, pr: Project) -> None:
        self._project = pr
        self.lbl_title.setText(f"Export: {pr.title}")
        self._load_templates()

    # -------- templates

    def _default_templates(self) -> List[ExportTemplate]:
        return [
            ExportTemplate(id="t_basic", name="Basic 2-column (Word / Explanation)", rtl=True, include_header=True, font_size=12),
            ExportTemplate(id="t_compact", name="Compact (smaller font)", rtl=True, include_header=True, font_size=10),
        ]

    def _load_templates(self) -> None:
        if self._templates_path.exists():
            try:
                data = json.loads(self._templates_path.read_text(encoding="utf-8"))
                self._templates = [ExportTemplate(**d) for d in data]
            except Exception:
                self._templates = self._default_templates()
        else:
            self._templates = self._default_templates()
            self._save_templates()

        self.lst.clear()
        for t in self._templates:
            self.lst.addItem(t.name)
        if self._templates:
            self.lst.setCurrentRow(0)

    def _save_templates(self) -> None:
        self._templates_path.write_text(
            json.dumps([asdict(t) for t in self._templates], ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    def _add_template(self) -> None:
        name, ok = QtWidgets.QInputDialog.getText(self, "New template", "Template name:")
        if not ok:
            return
        name = (name or "").strip()
        if not name:
            return
        tid = "t_" + str(abs(hash(f"{name}||{time.time()}")))
        self._templates.append(ExportTemplate(id=tid, name=name, rtl=True, include_header=True, font_size=12))
        self._save_templates()
        self._load_templates()
        self.lst.setCurrentRow(len(self._templates) - 1)

    def _render_preview(self, idx: int) -> None:
        if idx < 0 or idx >= len(self._templates):
            self.preview.setText("Pick a template to preview")
            return
        t = self._templates[idx]
        self.preview.setText(
            "Preview\n\n"
            f"• RTL: {'Yes' if t.rtl else 'No'}\n"
            f"• Header row: {'Yes' if t.include_header else 'No'}\n"
            f"• Font size: {t.font_size}\n\n"
            "Output will be a Word document with a 2-column table."
        )

    # -------- export

    def _do_export(self) -> None:
        if not self._project:
            return
        idx = self.lst.currentRow()
        if idx < 0 or idx >= len(self._templates):
            return
        tmpl = self._templates[idx]

        # Choose output path
        suggested = f"{self._project.title}.docx".replace("/", "-").replace("\\", "-")
        out_path, _ = QtWidgets.QFileDialog.getSaveFileName(
            self,
            "Export to Word",
            suggested,
            "Word Document (*.docx)",
        )
        if not out_path:
            return

        try:
            self._write_docx(Path(out_path), tmpl)
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Export failed", str(e))
            return

        # Register exported file
        self.ds.register_saved_export(
            export_id="x_" + str(abs(hash(f"{out_path}||{time.time()}"))),
            title=self._project.title,
            fmt="docx",
            source_project_id=self._project.id,
            out_path=str(out_path),
        )
        QtWidgets.QMessageBox.information(self, "Exported", f"Saved: {out_path}")

    def _write_docx(self, out_path: Path, tmpl: ExportTemplate) -> None:
        from docx import Document
        from docx.shared import Pt

        doc = Document()

        # Global font
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(int(tmpl.font_size))

        if tmpl.include_header:
            doc.add_heading(self._project.title, level=1)

        rows = self._project.rows or []
        table = doc.add_table(rows=len(rows) + (1 if tmpl.include_header else 0), cols=2)
        table.style = "Table Grid"

        r0 = 0
        if tmpl.include_header:
            hdr = table.rows[0].cells
            hdr[0].text = "Word"
            hdr[1].text = "Explanation"
            r0 = 1

        for i, row in enumerate(rows):
            w = (row.get("word") or "").strip()
            e = (row.get("explanation") or "").strip()
            cells = table.rows[r0 + i].cells
            cells[0].text = w
            cells[1].text = e

        doc.save(str(out_path))
